import json
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

# Define the file path to your Word document
file_path = "input.docx"

# Define the expected formatting for each section
title_format = {"font": "Arial", "size": 12, "bold": True}
abstract_format = {"font": "Arial", "size": 12}
body_format = {"font": "Calibri", "size": 11}
second_order_heading_format = {"font": "Calibri", "size": 12, "bold": True}
fifth_order_heading_format = {"font": "Calibri", "size": 11, "bold": True}

# Define the expected paper margins (in inches)
left_margin = 1
right_margin = 1
top_margin = 1
bottom_margin = 1

# Load the Word document using the docx library
document = docx.Document(file_path)

# Extract the formatting details for each section of the document
styles = {}
errors = []
for style in document.styles:
    if style.type == WD_STYLE_TYPE.PARAGRAPH:
        try:
            styles[style.name] = {
                "font": style.font.name,
                "size": style.font.size,
                "bold": style.font.bold,
                "italic": style.font.italic,
                "underline": style.font.underline,
                "highlight_color": style.font.highlight_color if style.font.highlight_color != WD_COLOR_INDEX.AUTO else None,
            }
        except KeyError:
            errors.append(f"{style.name} formatting is not defined in the expected formats")

title_style = styles.get("paper title")
if title_style is None:
    errors.append("Title formatting is not defined in the document")

abstract_style = styles.get("Abstract")
if abstract_style is None:
    errors.append("Abstract formatting is not defined in the document")

body_style = styles.get("Body Text")
if body_style is None:
    errors.append("Body formatting is not defined in the document")

second_order_heading_style = styles.get("Heading 2")
if second_order_heading_style is None:
    errors.append("Second order heading formatting is not defined in the document")

fifth_order_heading_style = styles.get("Heading 5")
if fifth_order_heading_style is None:
    errors.append("Fifth order heading formatting is not defined in the document")

# Check if the formatting of each section matches the expected formatting
if title_style != title_format:
    errors.append("Title formatting is incorrect")
if abstract_style != abstract_format:
    errors.append("Abstract formatting is incorrect")
if body_style != body_format:
    errors.append("Body formatting is incorrect")
if second_order_heading_style != second_order_heading_format:
    errors.append("Second order heading formatting is incorrect")
if fifth_order_heading_style != fifth_order_heading_format:
    errors.append("Fifth order heading formatting is incorrect")

print(errors)


#Reformatting the Document.
# Check if the formatting of each section matches the expected formatting
if title_style != title_format:
    errors.append("Title formatting is incorrect")
    # Modify the formatting of the title section
    for paragraph in document.paragraphs:
        if paragraph.style.name == "paper title":
            if title_format.get("font") is not None:
                paragraph.style.font.name = title_format.get("font")
            if title_format.get("size") is not None:
                paragraph.style.font.size = docx.shared.Pt(title_format.get("size"))
            if title_format.get("bold") is not None:
                paragraph.style.font.bold = title_format.get("bold")
else:
    # The formatting of the title section matches the expected formatting
    pass

if abstract_style != abstract_format:
    errors.append("Abstract formatting is incorrect")
    # Modify the formatting of the abstract section
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Abstract":
            if abstract_format.get("font") is not None:
                paragraph.style.font.name = abstract_format.get("font")
            if abstract_format.get("size") is not None:
                paragraph.style.font.size = docx.shared.Pt(abstract_format.get("size"))
else:
    # The formatting of the abstract section matches the expected formatting
    pass

if body_style != body_format:
    errors.append("Body formatting is incorrect")
    # Modify the formatting of the body text
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Body Text":
            if body_format.get("font") is not None:
                paragraph.style.font.name = body_format.get("font")
            if body_format.get("size") is not None:
                paragraph.style.font.size = docx.shared.Pt(body_format.get("size"))
else:
    # The formatting of the body text matches the expected formatting
    pass


if second_order_heading_style != second_order_heading_format:
    errors.append("second order heading style is incorrect")
    # Modify the formatting of the body text
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 2":
            if second_order_heading_format.get("font") is not None:
                paragraph.style.font.name = second_order_heading_format.get("font")
            if second_order_heading_format.get("size") is not None:
                paragraph.style.font.size = docx.shared.Pt(second_order_heading_format.get("size"))
else:
    # The formatting of the body text matches the expected formatting
    pass

if fifth_order_heading_style != fifth_order_heading_format:
    errors.append("second order heading style is incorrect")
    # Modify the formatting of the body text
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 5":
            if fifth_order_heading_format.get("font") is not None:
                paragraph.style.font.name = fifth_order_heading_format.get("font")
            if fifth_order_heading_format.get("size") is not None:
                paragraph.style.font.size = docx.shared.Pt(fifth_order_heading_format.get("size"))
else:
    # The formatting of the body text matches the expected formatting
    pass


# Save the modified document
document.save("output.docx")

